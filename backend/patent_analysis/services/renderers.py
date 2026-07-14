"""Pure report renderers.  They consume frozen State and never fetch data."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook


def render_docx(markdown: str, path: Path) -> None:
    document = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    document.save(path)


def render_xlsx(state, path: Path) -> None:
    workbook = Workbook()
    overview = workbook.active
    overview.title = "结论"
    overview.append(["项目", "值"])
    overview.append(["新颖性", state.novelty.overall if state.novelty else "uncertain"])
    overview.append(["创造性", state.inventiveness.overall if state.inventiveness else "uncertain"])
    overview.append(["商业价值置信度", state.commercial_value.overall_confidence if state.commercial_value else 0])

    features = workbook.create_sheet("特征")
    features.append(["ID", "类型", "限定", "文本"])
    for item in state.features:
        features.append([item.id, item.kind, item.limitation, item.text])

    evidence = workbook.create_sheet("证据")
    evidence.append(["ID", "文献", "特征", "定位", "引文", "已验证"])
    for item in state.evidence:
        evidence.append([item.id, item.document_id, ",".join(item.feature_ids), item.claim_number or item.paragraph_range or item.section, item.quoted_text, item.verified])
    workbook.save(path)
